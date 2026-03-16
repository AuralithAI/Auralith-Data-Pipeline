"""Tests for compound document extraction and ingestion."""

import io
import json
import struct
from pathlib import Path
from unittest import mock

import numpy as np
import pytest

# ── ModalitySegment / CompoundDocument ────────────────────────────────


class TestModalitySegment:
    """Basic tests for the ModalitySegment data class."""

    def test_text_segment(self):
        from auralith_pipeline.extraction.compound import ModalitySegment

        seg = ModalitySegment(modality="text", content_text="Hello world")
        assert seg.modality == "text"
        assert seg.content_text == "Hello world"
        assert seg.content_bytes is None

    def test_image_segment_with_bytes(self):
        from auralith_pipeline.extraction.compound import ModalitySegment

        seg = ModalitySegment(modality="image", content_bytes=b"\x89PNG")
        assert seg.modality == "image"
        assert seg.content_bytes == b"\x89PNG"
        assert seg.content_text == ""

    def test_metadata_default_empty(self):
        from auralith_pipeline.extraction.compound import ModalitySegment

        seg = ModalitySegment(modality="audio")
        assert seg.metadata == {}


class TestCompoundDocument:
    """Tests for CompoundDocument."""

    def test_modalities_present(self):
        from auralith_pipeline.extraction.compound import (
            CompoundDocument,
            ModalitySegment,
        )

        doc = CompoundDocument(
            source_path="test.docx",
            document_id="doc:123",
            segments=[
                ModalitySegment(modality="text", content_text="hello"),
                ModalitySegment(modality="image", content_bytes=b"\x89PNG"),
                ModalitySegment(modality="text", content_text="world"),
            ],
        )
        assert doc.modalities_present == {"text", "image"}
        assert doc.num_segments == 3


# ── Utility functions ─────────────────────────────────────────────────


class TestUtilities:
    """Tests for helper functions in the compound module."""

    def test_table_to_markdown(self):
        from auralith_pipeline.extraction.compound import _table_to_markdown

        rows = [["Name", "Age"], ["Alice", "30"], ["Bob", "25"]]
        md = _table_to_markdown(rows)
        assert "| Name | Age |" in md
        assert "| --- | --- |" in md
        assert "| Alice | 30 |" in md
        assert "| Bob | 25 |" in md

    def test_table_to_markdown_empty(self):
        from auralith_pipeline.extraction.compound import _table_to_markdown

        assert _table_to_markdown([]) == ""

    def test_table_to_markdown_pipe_escaping(self):
        from auralith_pipeline.extraction.compound import _table_to_markdown

        rows = [["A|B", "C"], ["D", "E"]]
        md = _table_to_markdown(rows)
        assert "A\\|B" in md

    def test_numpy_to_png_bytes_grayscale(self):
        from auralith_pipeline.extraction.compound import _numpy_to_png_bytes

        arr = np.zeros((32, 32), dtype=np.uint8)
        result = _numpy_to_png_bytes(arr)
        assert isinstance(result, bytes)
        assert result[:4] == b"\x89PNG"

    def test_numpy_to_png_bytes_rgb(self):
        from auralith_pipeline.extraction.compound import _numpy_to_png_bytes

        arr = np.zeros((32, 32, 3), dtype=np.uint8)
        result = _numpy_to_png_bytes(arr)
        assert isinstance(result, bytes)
        assert result[:4] == b"\x89PNG"


# ── CompoundDocumentExtractor ─────────────────────────────────────────


class TestCompoundDocumentExtractor:
    """Tests for the extractor."""

    def test_supported_extensions(self):
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        ext = CompoundDocumentExtractor()
        supported = ext.supported_extensions()
        assert ".docx" in supported
        assert ".xlsx" in supported
        assert ".pdf" in supported
        assert ".dcm" in supported
        assert ".tif" in supported
        assert ".mp4" in supported
        assert ".epub" in supported
        assert ".eml" in supported
        assert ".mbox" in supported
        assert ".rtf" in supported
        assert ".mhtml" in supported
        assert ".tex" in supported
        assert ".h5" in supported
        assert ".fits" in supported
        assert ".zarr" in supported
        assert ".warc" in supported
        assert ".glb" in supported
        assert ".gltf" in supported
        assert ".cbz" in supported
        assert ".cbr" in supported
        assert ".msg" in supported

    def test_unsupported_extension_raises(self, tmp_path: Path):
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        ext = CompoundDocumentExtractor()
        fake = tmp_path / "test.unsupported_ext_xyz"
        fake.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported compound file type"):
            ext.extract(fake)

    def test_extract_docx_text(self, tmp_path: Path):
        """Test DOCX extraction with mocked python-docx."""
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False)

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("Second paragraph")
            docx_path = tmp_path / "test.docx"
            doc.save(str(docx_path))
        except ImportError:
            pytest.skip("python-docx not installed")

        result = extractor.extract(docx_path)
        assert result.num_segments >= 1
        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        combined_text = " ".join(s.content_text for s in text_segs)
        assert "First paragraph" in combined_text
        assert "Second paragraph" in combined_text

    def test_extract_docx_with_table(self, tmp_path: Path):
        """DOCX with a table should produce a table segment in Markdown."""
        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Some text")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Alpha"
            table.cell(1, 1).text = "100"
            docx_path = tmp_path / "with_table.docx"
            doc.save(str(docx_path))
        except ImportError:
            pytest.skip("python-docx not installed")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False, extract_tables=True)
        result = extractor.extract(docx_path)

        table_segs = [s for s in result.segments if s.metadata.get("type") == "docx_table"]
        assert len(table_segs) >= 1
        assert "Alpha" in table_segs[0].content_text
        assert "| Name |" in table_segs[0].content_text

    def test_extract_xlsx(self, tmp_path: Path):
        """XLSX extraction should produce sheet text segments."""
        try:
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws.title = "Sheet1"
            ws.append(["Name", "Score"])
            ws.append(["Alice", 95])
            ws.append(["Bob", 87])
            xlsx_path = tmp_path / "test.xlsx"
            wb.save(str(xlsx_path))
        except ImportError:
            pytest.skip("openpyxl not installed")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False)
        result = extractor.extract(xlsx_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        combined = " ".join(s.content_text for s in text_segs)
        assert "Alice" in combined
        assert "Bob" in combined

    def test_extract_hgt_elevation(self, tmp_path: Path):
        """HGT elevation file should produce image tiles."""
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        # Create a small synthetic HGT file (3x3 = 9 values, but HGT must
        # be square with side = sqrt(num_values). Use 1201x1201 = SRTM3 is
        # too large for tests, so use a small 4x4.
        side = 64
        num_values = side * side
        elevations = list(range(num_values))
        raw = struct.pack(f">{num_values}h", *elevations)
        hgt_path = tmp_path / "test.hgt"
        hgt_path.write_bytes(raw)

        extractor = CompoundDocumentExtractor(geotiff_tile_size=32)
        result = extractor.extract(hgt_path)

        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(image_segs) >= 1
        assert all(s.content_bytes is not None for s in image_segs)
        # Verify the bytes start with PNG header
        assert image_segs[0].content_bytes[:4] == b"\x89PNG"


# ── CompoundDocumentSource ────────────────────────────────────────────


class TestCompoundDocumentSource:
    """Tests for the data source wrapper."""

    def test_source_name(self, tmp_path: Path):
        from auralith_pipeline.sources.compound import CompoundDocumentSource

        source = CompoundDocumentSource(root_dir=tmp_path)
        assert "compound:" in source.name

    def test_empty_directory(self, tmp_path: Path):
        from auralith_pipeline.sources.compound import CompoundDocumentSource

        source = CompoundDocumentSource(root_dir=tmp_path)
        assert len(source) == 0
        assert list(source) == []

    def test_source_yields_data_samples(self, tmp_path: Path):
        """End-to-end: create a DOCX, ingest via CompoundDocumentSource."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Test compound document source integration")
        docx_path = tmp_path / "sample.docx"
        doc.save(str(docx_path))

        from auralith_pipeline.sources.compound import CompoundDocumentSource

        source = CompoundDocumentSource(
            root_dir=tmp_path,
            extract_images=False,
            extract_tables=True,
        )
        assert len(source) == 1

        samples = list(source)
        assert len(samples) >= 1
        assert samples[0].modality == "text"
        assert "compound document source" in samples[0].content

    def test_max_samples_respected(self, tmp_path: Path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        for i in range(3):
            doc = Document()
            doc.add_paragraph(f"Document {i} paragraph one")
            doc.add_paragraph(f"Document {i} paragraph two")
            doc.save(str(tmp_path / f"doc_{i}.docx"))

        from auralith_pipeline.sources.compound import CompoundDocumentSource

        source = CompoundDocumentSource(root_dir=tmp_path, max_samples=2, extract_images=False)
        samples = list(source)
        assert len(samples) <= 2

    def test_binary_assets_staged_to_disk(self, tmp_path: Path):
        """Image segments should be staged to disk with binary_path in metadata."""
        from auralith_pipeline.extraction.compound import (
            ModalitySegment,
        )
        from auralith_pipeline.sources.compound import CompoundDocumentSource

        source = CompoundDocumentSource(root_dir=tmp_path, extract_images=True)

        png_bytes = _make_minimal_png()
        seg = ModalitySegment(
            modality="image",
            content_bytes=png_bytes,
            metadata={"source": "test.docx", "type": "docx_embedded_image"},
        )

        sample = source._segment_to_sample(seg, "test_doc_id")
        assert sample is not None
        assert sample.modality == "image"
        assert "binary_path" in sample.metadata
        assert Path(sample.metadata["binary_path"]).exists()


# ── COMPOUND_EXTS coverage ────────────────────────────────────────────


class TestCompoundExts:
    """Ensure COMPOUND_EXTS is consistent."""

    def test_compound_exts_is_frozenset(self):
        from auralith_pipeline.extraction.compound import COMPOUND_EXTS

        assert isinstance(COMPOUND_EXTS, frozenset)

    def test_key_formats_present(self):
        from auralith_pipeline.extraction.compound import COMPOUND_EXTS

        expected = [
            ".docx",
            ".xlsx",
            ".pdf",
            ".pptx",
            ".dcm",
            ".nii",
            ".tif",
            ".mp4",
            ".epub",
            ".eml",
            ".mbox",
            ".rtf",
            ".mhtml",
            ".mht",
            ".tex",
            ".h5",
            ".hdf5",
            ".fits",
            ".zarr",
            ".warc",
            ".glb",
            ".gltf",
            ".cbz",
            ".cbr",
            ".msg",
        ]
        for ext in expected:
            assert ext in COMPOUND_EXTS, f"{ext} missing from COMPOUND_EXTS"

    def test_no_overlap_with_code_exts(self):
        from auralith_pipeline.extraction.compound import COMPOUND_EXTS
        from auralith_pipeline.utils.file_types import CODE_EXTS

        overlap = COMPOUND_EXTS & CODE_EXTS
        assert overlap == frozenset(), f"COMPOUND_EXTS overlaps with CODE_EXTS: {overlap}"


# ── EPUB ──────────────────────────────────────────────────────


class TestExtractEpub:
    """Tests for EPUB extraction."""

    def test_extract_epub_text(self, tmp_path: Path):
        """EPUB extraction should produce text segments from XHTML content."""
        import zipfile

        epub_path = tmp_path / "test.epub"
        with zipfile.ZipFile(str(epub_path), "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr(
                "OEBPS/chapter1.xhtml",
                "<html><body><p>This is a test chapter with enough content to pass the length check.</p></body></html>",
            )
            zf.writestr(
                "OEBPS/chapter2.xhtml",
                "<html><body><h1>Chapter Two</h1><p>Another chapter with some text for testing purposes.</p></body></html>",
            )

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False)
        result = extractor.extract(epub_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 2
        combined = " ".join(s.content_text for s in text_segs)
        assert "test chapter" in combined
        assert "Chapter Two" in combined

    def test_extract_epub_images(self, tmp_path: Path):
        """EPUB with embedded images should produce image segments."""
        import zipfile

        png_bytes = _make_minimal_png()
        epub_path = tmp_path / "with_images.epub"
        with zipfile.ZipFile(str(epub_path), "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr(
                "OEBPS/chapter1.xhtml",
                "<html><body><p>Text with enough characters to pass threshold limit easily here.</p></body></html>",
            )
            zf.writestr("OEBPS/images/cover.png", png_bytes)

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=True)
        result = extractor.extract(epub_path)

        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(image_segs) >= 1
        assert image_segs[0].metadata["type"] == "epub_image"

    def test_extract_epub_skips_scripts(self, tmp_path: Path):
        """Script tags should be stripped from EPUB text."""
        import zipfile

        epub_path = tmp_path / "scripts.epub"
        with zipfile.ZipFile(str(epub_path), "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr(
                "OEBPS/chapter.xhtml",
                "<html><body><p>Visible text that is long enough to pass the threshold.</p><script>var evil = true;</script></body></html>",
            )

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False)
        result = extractor.extract(epub_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        assert "evil" not in text_segs[0].content_text
        assert "Visible text" in text_segs[0].content_text


# ── EML / Email ──────────────────────────────────────────────


class TestExtractEmail:
    """Tests for EML and mbox extraction."""

    def test_extract_eml_plain_text(self, tmp_path: Path):
        """EML file with plain text body."""
        eml_content = (
            "From: sender@test.com\r\n"
            "To: receiver@test.com\r\n"
            "Subject: Test Email\r\n"
            "Content-Type: text/plain; charset=utf-8\r\n"
            "\r\n"
            "Hello, this is the email body."
        )
        eml_path = tmp_path / "test.eml"
        eml_path.write_text(eml_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(eml_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        combined = " ".join(s.content_text for s in text_segs)
        assert "Test Email" in combined
        assert "email body" in combined

    def test_extract_eml_multipart_with_image(self, tmp_path: Path):
        """EML with multipart content and image attachment."""
        from email.mime.image import MIMEImage
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText

        msg = MIMEMultipart()
        msg["Subject"] = "Multipart Test"
        msg["From"] = "sender@test.com"
        msg.attach(MIMEText("Body text here", "plain"))
        msg.attach(MIMEImage(_make_minimal_png(), "png", name="test.png"))

        eml_path = tmp_path / "multipart.eml"
        eml_path.write_bytes(msg.as_bytes())

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=True)
        result = extractor.extract(eml_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(text_segs) >= 1
        assert len(image_segs) >= 1
        assert "Body text" in text_segs[0].content_text

    def test_extract_mbox(self, tmp_path: Path):
        """Mbox file with multiple messages."""
        mbox_content = (
            "From sender@test.com Mon Jan  1 00:00:00 2024\r\n"
            "From: sender@test.com\r\n"
            "Subject: First message\r\n"
            "Content-Type: text/plain\r\n"
            "\r\n"
            "First message body\r\n"
            "\r\n"
            "From sender@test.com Mon Jan  2 00:00:00 2024\r\n"
            "From: sender@test.com\r\n"
            "Subject: Second message\r\n"
            "Content-Type: text/plain\r\n"
            "\r\n"
            "Second message body\r\n"
        )
        mbox_path = tmp_path / "test.mbox"
        mbox_path.write_text(mbox_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(mbox_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 2
        combined = " ".join(s.content_text for s in text_segs)
        assert "First message" in combined
        assert "Second message" in combined


# ── RTF ───────────────────────────────────────────────────────


class TestExtractRtf:
    """Tests for RTF extraction."""

    def test_extract_rtf_text(self, tmp_path: Path):
        """RTF file should produce text."""
        rtf_content = r"{\rtf1\ansi Hello RTF World!}"
        rtf_path = tmp_path / "test.rtf"
        rtf_path.write_text(rtf_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(rtf_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        assert "Hello RTF World" in text_segs[0].content_text

    def test_extract_rtf_metadata(self, tmp_path: Path):
        """RTF segment should have correct metadata type."""
        rtf_content = r"{\rtf1\ansi Some text content for metadata check.}"
        rtf_path = tmp_path / "meta.rtf"
        rtf_path.write_text(rtf_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(rtf_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        assert text_segs[0].metadata["type"] == "rtf_text"


# ── MHTML ─────────────────────────────────────────────────────


class TestExtractMhtml:
    """Tests for MHTML web archive extraction."""

    def test_extract_mhtml_text(self, tmp_path: Path):
        """MHTML file with HTML content should produce text."""
        mhtml_content = (
            "MIME-Version: 1.0\r\n"
            "Content-Type: text/html; charset=utf-8\r\n"
            "\r\n"
            "<html><body><h1>Archived Page</h1><p>This is archived content.</p></body></html>"
        )
        mhtml_path = tmp_path / "test.mhtml"
        mhtml_path.write_text(mhtml_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(mhtml_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        assert "Archived Page" in text_segs[0].content_text
        assert "archived content" in text_segs[0].content_text

    def test_extract_mht_extension(self, tmp_path: Path):
        """The .mht extension should use the same handler."""
        mht_content = (
            "MIME-Version: 1.0\r\n"
            "Content-Type: text/plain; charset=utf-8\r\n"
            "\r\n"
            "Plain text MHT file"
        )
        mht_path = tmp_path / "test.mht"
        mht_path.write_text(mht_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(mht_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        assert "Plain text MHT" in text_segs[0].content_text


# ── LaTeX ─────────────────────────────────────────────────────


class TestExtractLatex:
    """Tests for LaTeX extraction."""

    def test_extract_latex_text(self, tmp_path: Path):
        """LaTeX file should produce cleaned text."""
        tex_content = r"""\documentclass{article}
\usepackage{amsmath}

\begin{document}

\section{Introduction}

This is an introduction to our paper with enough text for the threshold.

\subsection{Background}

Some background information about the topic we are studying here.

\end{document}
"""
        tex_path = tmp_path / "paper.tex"
        tex_path.write_text(tex_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False)
        result = extractor.extract(tex_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        text = text_segs[0].content_text
        assert "Introduction" in text
        assert "Background" in text
        assert "\\documentclass" not in text

    def test_extract_latex_figures(self, tmp_path: Path):
        """LaTeX with includegraphics should extract referenced images."""
        png_bytes = _make_minimal_png()
        (tmp_path / "fig1.png").write_bytes(png_bytes)

        tex_content = r"""\documentclass{article}
\begin{document}
This is a paper with a figure and enough text to pass the length threshold.
\includegraphics[width=0.5\textwidth]{fig1}
\end{document}
"""
        tex_path = tmp_path / "paper.tex"
        tex_path.write_text(tex_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=True)
        result = extractor.extract(tex_path)

        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(image_segs) >= 1
        assert image_segs[0].metadata["type"] == "latex_figure"

    def test_extract_latex_math_replaced(self, tmp_path: Path):
        """Inline math should be replaced with [math] placeholder."""
        tex_content = r"""\documentclass{article}
\begin{document}
The equation $E = mc^2$ is famous enough text to pass the threshold for extraction.
\end{document}
"""
        tex_path = tmp_path / "math.tex"
        tex_path.write_text(tex_content, encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False)
        result = extractor.extract(tex_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        assert "[math]" in text_segs[0].content_text


# ── HDF5 ──────────────────────────────────────────────────────


class TestExtractHdf5:
    """Tests for HDF5 extraction (mocked to avoid h5py dependency)."""

    def test_extract_hdf5_2d_dataset(self, tmp_path: Path):
        """HDF5 with a 2-D numerical dataset should produce image segment."""
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        mock_dataset = mock.MagicMock()
        mock_dataset.shape = (32, 32)
        mock_dataset.dtype = np.dtype("float32")
        mock_dataset.__getitem__ = mock.MagicMock(
            return_value=np.random.rand(32, 32).astype(np.float32)
        )
        mock_dataset.attrs = {}

        mock_file = mock.MagicMock()
        mock_file.__enter__ = mock.MagicMock(return_value=mock_file)
        mock_file.__exit__ = mock.MagicMock(return_value=False)

        def mock_visititems(fn):
            fn("dataset1", mock_dataset)

        mock_file.visititems = mock_visititems

        h5_path = tmp_path / "test.h5"
        h5_path.write_bytes(b"fake")

        with mock.patch.dict("sys.modules", {"h5py": mock.MagicMock()}):
            import sys

            mock_h5py = sys.modules["h5py"]
            mock_h5py.File.return_value = mock_file
            mock_h5py.Dataset = type(mock_dataset)

            extractor = CompoundDocumentExtractor()
            result = extractor.extract(h5_path)

        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(image_segs) >= 1


# ── FITS ──────────────────────────────────────────────────────


class TestExtractFits:
    """Tests for FITS extraction (mocked to avoid astropy dependency)."""

    def test_extract_fits_image(self, tmp_path: Path):
        """FITS with image data should produce image segments."""
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        mock_header = mock.MagicMock()
        mock_header.__iter__ = mock.MagicMock(return_value=iter(["NAXIS", "NAXIS1"]))
        mock_header.__getitem__ = mock.MagicMock(return_value="512")
        mock_header.comments = {"NAXIS": "number of axes", "NAXIS1": "length of axis 1"}

        mock_hdu = mock.MagicMock()
        mock_hdu.header = mock_header
        mock_hdu.data = np.random.rand(32, 32).astype(np.float32)

        # The context manager must return a real iterable of HDUs.
        hdu_list = [mock_hdu]

        mock_hdul = mock.MagicMock()
        mock_hdul.__enter__ = mock.MagicMock(return_value=hdu_list)
        mock_hdul.__exit__ = mock.MagicMock(return_value=False)

        fits_path = tmp_path / "test.fits"
        fits_path.write_bytes(b"fake fits")

        # Wire up mock so `from astropy.io import fits` resolves correctly
        mock_fits_module = mock.MagicMock()
        mock_fits_module.open.return_value = mock_hdul
        mock_astropy_io = mock.MagicMock()
        mock_astropy_io.fits = mock_fits_module

        with mock.patch.dict(
            "sys.modules",
            {
                "astropy": mock.MagicMock(),
                "astropy.io": mock_astropy_io,
                "astropy.io.fits": mock_fits_module,
            },
        ):
            extractor = CompoundDocumentExtractor()
            result = extractor.extract(fits_path)

        assert result.num_segments >= 1


# ── Zarr ──────────────────────────────────────────────────────


class TestExtractZarr:
    """Tests for Zarr extraction (mocked to avoid zarr dependency)."""

    def test_extract_zarr_2d_array(self, tmp_path: Path):
        """Zarr with 2-D array should produce image segment."""
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        mock_array = mock.MagicMock()
        mock_array.shape = (32, 32)
        mock_array.dtype = np.dtype("float32")
        mock_array.__getitem__ = mock.MagicMock(
            return_value=np.random.rand(32, 32).astype(np.float32)
        )

        mock_store = mock.MagicMock()
        mock_store.attrs = {}
        mock_store.arrays.return_value = [("data", mock_array)]
        del mock_store.visititems

        zarr_path = tmp_path / "test.zarr"
        zarr_path.mkdir()

        mock_zarr = mock.MagicMock()
        mock_zarr.open.return_value = mock_store

        with mock.patch.dict("sys.modules", {"zarr": mock_zarr}):
            extractor = CompoundDocumentExtractor()
            result = extractor.extract(zarr_path)

        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(image_segs) >= 1


# ── WARC ──────────────────────────────────────────────────────


class TestExtractWarc:
    """Tests for WARC extraction (mocked to avoid warcio dependency)."""

    def test_extract_warc_html(self, tmp_path: Path):
        """WARC with HTML response should produce text segment."""
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        mock_http_headers = mock.MagicMock()
        mock_http_headers.get_header.return_value = "text/html; charset=utf-8"

        mock_rec_headers = mock.MagicMock()
        mock_rec_headers.get_header.return_value = "http://example.com/page"

        html = "<html><body><p>This is web archive content that is long enough for extraction threshold.</p></body></html>"
        mock_stream = io.BytesIO(html.encode("utf-8"))

        mock_record = mock.MagicMock()
        mock_record.rec_type = "response"
        mock_record.http_headers = mock_http_headers
        mock_record.rec_headers = mock_rec_headers
        mock_record.content_stream.return_value = mock_stream

        mock_warcio = mock.MagicMock()
        mock_warcio.archiveiterator.ArchiveIterator.return_value = [mock_record]

        warc_path = tmp_path / "test.warc"
        warc_path.write_bytes(b"fake warc")

        with mock.patch.dict(
            "sys.modules",
            {
                "warcio": mock_warcio,
                "warcio.archiveiterator": mock_warcio.archiveiterator,
            },
        ):
            extractor = CompoundDocumentExtractor()
            result = extractor.extract(warc_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        assert "web archive content" in text_segs[0].content_text


# ── glTF / GLB ───────────────────────────────────────────────


class TestExtractGltf:
    """Tests for glTF/GLB extraction."""

    def test_extract_gltf_metadata(self, tmp_path: Path):
        """glTF file should produce metadata text segment."""
        gltf_data = {
            "asset": {"generator": "TestGen", "version": "2.0"},
            "scenes": [{"name": "MainScene"}],
            "nodes": [{"name": "Cube"}, {"name": "Camera"}],
            "materials": [{"name": "DefaultMat"}],
            "meshes": [{"name": "CubeMesh"}],
        }
        gltf_path = tmp_path / "model.gltf"
        gltf_path.write_text(json.dumps(gltf_data), encoding="utf-8")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False)
        result = extractor.extract(gltf_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        text = text_segs[0].content_text
        assert "TestGen" in text
        assert "Cube" in text
        assert "DefaultMat" in text

    def test_extract_glb_header_validation(self, tmp_path: Path):
        """GLB with invalid header should not crash."""
        glb_path = tmp_path / "invalid.glb"
        glb_path.write_bytes(b"not a valid glb file at all")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(glb_path)
        assert result is not None

    def test_extract_glb_valid(self, tmp_path: Path):
        """GLB with valid header and JSON chunk should extract metadata."""
        gltf_json = json.dumps(
            {
                "asset": {"version": "2.0", "generator": "GLBTest"},
                "meshes": [{"name": "TestMesh"}],
            }
        ).encode("utf-8")
        while len(gltf_json) % 4 != 0:
            gltf_json += b" "

        header = struct.pack("<III", 0x46546C67, 2, 12 + 8 + len(gltf_json))
        chunk_header = struct.pack("<II", len(gltf_json), 0x4E4F534A)

        glb_path = tmp_path / "valid.glb"
        glb_path.write_bytes(header + chunk_header + gltf_json)

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor(extract_images=False)
        result = extractor.extract(glb_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        assert "GLBTest" in text_segs[0].content_text


# ── Comic Book Archives ──────────────────────────────────────


class TestExtractComicBook:
    """Tests for CBZ/CBR extraction."""

    def test_extract_cbz(self, tmp_path: Path):
        """CBZ (ZIP of images) should extract page images."""
        import zipfile

        png_bytes = _make_minimal_png()
        cbz_path = tmp_path / "comic.cbz"
        with zipfile.ZipFile(str(cbz_path), "w") as zf:
            zf.writestr("page_001.png", png_bytes)
            zf.writestr("page_002.jpg", png_bytes)
            zf.writestr("page_003.png", png_bytes)

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(cbz_path)

        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(image_segs) == 3
        assert all(s.metadata["type"] == "comic_page" for s in image_segs)
        # Should be in sorted order
        assert image_segs[0].metadata["page_index"] == 0
        assert image_segs[2].metadata["page_index"] == 2

    def test_extract_cbz_skips_macosx_folder(self, tmp_path: Path):
        """CBZ should skip __MACOSX metadata entries."""
        import zipfile

        png_bytes = _make_minimal_png()
        cbz_path = tmp_path / "mac_comic.cbz"
        with zipfile.ZipFile(str(cbz_path), "w") as zf:
            zf.writestr("page_001.png", png_bytes)
            zf.writestr("__MACOSX/._page_001.png", b"metadata")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(cbz_path)

        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(image_segs) == 1

    def test_extract_cbr_without_rarfile(self, tmp_path: Path):
        """CBR without rarfile installed should gracefully handle error."""
        cbr_path = tmp_path / "comic.cbr"
        cbr_path.write_bytes(b"Rar!\x1a\x07\x00fake")

        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        extractor = CompoundDocumentExtractor()
        result = extractor.extract(cbr_path)
        assert result is not None


# ── Outlook MSG ──────────────────────────────────────────────


class TestExtractMsg:
    """Tests for Outlook .msg extraction (mocked)."""

    def test_extract_msg_body(self, tmp_path: Path):
        """MSG file should produce text from subject + body."""
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        mock_msg = mock.MagicMock()
        mock_msg.subject = "Important Meeting"
        mock_msg.body = "Please attend the meeting at 3 PM."
        mock_msg.date = "2024-01-15"
        mock_msg.attachments = []

        mock_extract_msg = mock.MagicMock()
        mock_extract_msg.Message.return_value = mock_msg

        msg_path = tmp_path / "email.msg"
        msg_path.write_bytes(b"fake msg")

        with mock.patch.dict("sys.modules", {"extract_msg": mock_extract_msg}):
            extractor = CompoundDocumentExtractor()
            result = extractor.extract(msg_path)

        text_segs = [s for s in result.segments if s.modality == "text"]
        assert len(text_segs) >= 1
        combined = text_segs[0].content_text
        assert "Important Meeting" in combined
        assert "meeting at 3 PM" in combined

    def test_extract_msg_with_image_attachment(self, tmp_path: Path):
        """MSG with image attachment should produce image segment."""
        from auralith_pipeline.extraction.compound import CompoundDocumentExtractor

        mock_att = mock.MagicMock()
        mock_att.mimetype = "image/png"
        mock_att.data = _make_minimal_png()
        mock_att.longFilename = "photo.png"

        mock_msg = mock.MagicMock()
        mock_msg.subject = "Photo"
        mock_msg.body = "See attached."
        mock_msg.date = None
        mock_msg.attachments = [mock_att]

        mock_extract_msg = mock.MagicMock()
        mock_extract_msg.Message.return_value = mock_msg

        msg_path = tmp_path / "photo.msg"
        msg_path.write_bytes(b"fake msg")

        with mock.patch.dict("sys.modules", {"extract_msg": mock_extract_msg}):
            extractor = CompoundDocumentExtractor(extract_images=True)
            result = extractor.extract(msg_path)

        image_segs = [s for s in result.segments if s.modality == "image"]
        assert len(image_segs) >= 1


# ── Utility: _normalise_to_uint8 ─────────────────────────────────────


class TestNormaliseToUint8:
    """Tests for the _normalise_to_uint8 helper."""

    def test_normalise_range(self):
        from auralith_pipeline.extraction.compound import _normalise_to_uint8

        arr = np.array([0.0, 0.5, 1.0], dtype=np.float32)
        result = _normalise_to_uint8(arr)
        assert result.dtype == np.uint8
        assert result[0] == 0
        assert result[2] == 255

    def test_normalise_constant(self):
        from auralith_pipeline.extraction.compound import _normalise_to_uint8

        arr = np.full((4, 4), 42.0, dtype=np.float32)
        result = _normalise_to_uint8(arr)
        assert result.dtype == np.uint8
        assert np.all(result == 0)


# ── Helpers ───────────────────────────────────────────────────────────


def _make_minimal_png() -> bytes:
    """Create a minimal valid 1x1 PNG for testing."""
    from PIL import Image

    buf = io.BytesIO()
    img = Image.new("RGB", (1, 1), color=(255, 0, 0))
    img.save(buf, format="PNG")
    return buf.getvalue()
